import math
import docx
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

from haggis.files.docx import list_number
import re

from app.models.question import Question
from typing import List, Dict


class HTMLHelper(object):
    """ Translates some html into word runs. """

    def __init__(self):
        self.get_tags = re.compile(
            "(<[a-z,A-Z]+>)(.*?)(</[a-z,A-Z]+>)", flags=re.DOTALL)

    def html_to_run_map(self, html_fragment):
        """ Breaks an html fragment into a run map """
        ptr = 0
        run_map = []
        for match in self.get_tags.finditer(html_fragment):
            if match.start() > ptr:
                text = html_fragment[ptr:match.start()]
                if len(text) > 0:
                    run_map.append((text, "plain_text"))
            run_map.append((match.group(2), match.group(1)))
            ptr = match.end()
        if ptr <= len(html_fragment) - 1:
            run_map.append((html_fragment[ptr:], "plain_text"))
        return run_map

    def insert_runs_from_html_map(self, paragraph, run_map):
        """ Inserts some runs into a paragraph object. """
        for run_item in run_map:
            run = paragraph.add_run(run_item[0])
            if run_item[1] == "<b>":
                run.bold = True
            elif run_item[1] == "<i>":
                run.italic = True


def parse_questions(questions: Dict[str, List[Question]], doc: docx.Document) -> None:
    """Parse sorted questions into a docx document

    Args:
        sorted_questions (Dict[str, List[Question]]): Dict of questions with section as key
        doc (docx.Document): docx document to write to
    """

    # Parse MCQ questions
    parse_mcq_question(questions["A"], doc)

    # Add MCQ answers
    add_mcq_answers(questions["A"], doc)

    # Parse structured questions
    parse_structured_question(questions["B"], doc)
    parse_structured_question(questions["C"], doc)

    # Add structured answers
    add_structured_answers(questions["B"], doc)
    add_structured_answers(questions["C"], doc)


def parse_mcq_question(questions: List[Question], doc: docx.Document) -> None:
    """Parse a multiple choice question into a docx document

    Args:
        qn_no (int): Question number
        question (Question): MCQ Question to parse
        doc (docx.Document): docx document to write to
    """

    # Return if no questions
    if len(questions) == 0:
        return

    # Add "root" paragraph
    prev_root = doc.add_paragraph()
    prev_root.paragraph_format.left_indent = docx.shared.Cm(0.5)
    prev_root.style = doc.styles["Question"]
    prev_root.paragraph_format.keep_with_next = True

    # Loop through and add questions
    for i, question in enumerate(questions):
        # Add question number
        if i == 0:
            root = prev_root
        else:
            root = doc.add_paragraph()
            root.paragraph_format.left_indent = docx.shared.Cm(0.5)
            root.paragraph_format.space_before = docx.shared.Pt(20)
            root.style = doc.styles["Question"]
            root.paragraph_format.keep_with_next = True

        # Add question
        for j, block in enumerate(question.content["blocks"]):
            # Create new paragraph if not first block
            if j != 0:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = docx.shared.Cm(0.5)
                para.paragraph_format.space_before = docx.shared.Pt(8)
                para.style = doc.styles["Question"]
                para.paragraph_format.keep_with_next = True
            else:
                para = root

            # Override space_after to 0 if last block, for answer place
            if j == len(question.content["blocks"]) - 1:
                para.paragraph_format.space_after = docx.shared.Pt(0)

            if block["type"] == "paragraph":
                add_text(block, para)
            elif block["type"] == "image":
                add_image(block, para)
            elif block["type"] == "table":
                add_table(block, doc)

        # Add paragraph for answer place
        ans_para = doc.add_paragraph()
        ans_para.paragraph_format.left_indent = docx.shared.Cm(0)
        ans_para.paragraph_format.space_before = docx.shared.Pt(0)
        ans_para.style = doc.styles["Question"]

        # Add text for answer place
        ans_para.add_run("(        )")
        ans_para.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT

        # Make paragraph a list to have numbering
        if i != 0:
            list_number(doc, root, prev=prev_root)
        else:
            list_number(doc, root)

    # Add page break
    doc.add_page_break()


def parse_structured_question(questions: List[Question], doc: docx.Document) -> None:
    """Parse a structured question into a docx document

    Args:
        qn_no (int): Question number
        question (Question): Structured Question to parse
        doc (docx.Document): docx document to write to
    """

    # Return if no questions
    if len(questions) == 0:
        return

    pass


def add_mcq_answers(questions: List[Question], doc: docx.Document) -> None:
    """Add MCQ answers to a docx document

    Args:
        questions (List[Question]): List of questions
        doc (docx.Document): docx document to write to
    """

    # Return if no questions
    if len(questions) == 0:
        return

    # Add MCQ heading
    p = doc.add_paragraph("MCQ Answers")
    p.style = doc.styles["Answer Heading"]
    p.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    # Insert line
    insertHR(p)

    # Create table
    answer_table = doc.add_table(rows=0, cols=4)
    answer_table.style = "Table Grid"

    # Align table to centre
    answer_table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER

    # Calculate total rows to be in table
    num_questions = len(questions)
    total_rows = math.ceil(num_questions / 2)

    # Populate tables with answers
    for i in range(total_rows):
        # Add "left" column
        row_cells = answer_table.add_row().cells
        row_cells[0].text = str(i + 1)
        row_cells[1].text = sanitise_mcq_answer(
            questions[i].answer["blocks"][0]["data"]["text"])

        # Add "right" column if total qns are not odd
        if (i + total_rows) < num_questions:
            row_cells[2].text = str(i + 1 + total_rows)
            row_cells[3].text = sanitise_mcq_answer(
                questions[i].answer["blocks"][0]["data"]["text"])

    # Style the table
    for row in answer_table.rows:
        for i, cell in enumerate(row.cells):
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = docx.shared.Pt(20)
                    paragraph.paragraph_format.alignment = 1

                    # Set font to bold if it is the first or third cell in the row
                    # As they are the question numbers
                    if i == 0 or i == 2:
                        run.font.bold = True

    # Add page break
    doc.add_page_break()


def add_structured_answers(questions: List[Question], doc: docx.Document) -> None:
    """Add structured answers to a docx document

    Args:
        questions (List[Question]): List of questions
        doc (docx.Document): docx document to write to
    """

    # Return if no questions
    if len(questions) == 0:
        return

    pass


def add_text(block: Dict[str, str], paragraph: docx.text.paragraph.Paragraph) -> None:
    """Add text to a paragraph

    Args:
        block (Dict[str, str]): EditorJS block to add
        paragraph (docx.text.paragraph.Paragraph): Paragraph to add to
    """

    # Get alignment, default to left
    try:
        alignment = block["data"]["alignment"]
    except KeyError:
        alignment = "left"

    # Add text to paragraph
    add_text_from_str(block["data"]["text"], paragraph, alignment=alignment)


def add_text_from_str(text: str, paragraph: docx.text.paragraph.Paragraph, alignment: str = "left") -> None:
    """Add text to a paragraph

    Args:
        block (Dict[str, str]): EditorJS block to add
        paragraph (docx.text.paragraph.Paragraph): Paragraph to add to
    """

    # Initialise HTML helper
    html_helper = HTMLHelper()

    # Process text
    text = process_text(text)
    run_map = html_helper.html_to_run_map(text)

    # Add text to paragraph and add new line if not a table
    html_helper.insert_runs_from_html_map(paragraph, run_map)

    # Set alignment
    if alignment == "left":
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
    elif alignment == "right":
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
    elif alignment == "center":
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER


def add_image(block: Dict[str, str], paragraph: docx.text.paragraph.Paragraph) -> None:
    """Add image to a paragraph

    NOTE: The image is not aligned to the centre, but rather
    padded with spaces to the left to make it not touch the left.

    Args:
        block (Dict[str, str]): EditorJS block to add
        paragraph (docx.text.paragraph.Paragraph): Paragraph to add to
    """

    # Get image path
    image_path = block["data"]["file"]["url"]
    rel_path = "./app" + image_path

    # Add image
    image = paragraph.add_run()
    image.add_picture(rel_path, width=docx.shared.Cm(13))

    # Align image to centre
    paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER


def add_table(block, doc) -> None:
    # Remove last paragraph to prevent newline above table
    delete_paragraph(doc.paragraphs[-1])

    # Get the value of withHeadings
    with_headings = block["data"]["withHeadings"]

    # Get table data
    table_data = block["data"]["content"]

    # Create table and set style
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = "Table Grid"

    # Force table to be on one page
    table.allow_autofit = False

    # Add indent to table
    tbl_pr = table._tblPr
    e = OxmlElement("w:tblInd")
    e.set(qn("w:w"), "350")  # TODO: Find the correct value
    e.set(qn("w:type"), "dxa")
    tbl_pr.append(e)

    # Set table height to 0.74cm
    for row in table.rows:
        row.height = docx.shared.Cm(0.74)

    # Add table data
    for i, row in enumerate(table_data):
        for j, cell in enumerate(row):
            # Get current cell
            curr_cell = table.cell(i, j)

            # Align cell to centre
            curr_cell.vertical_alignment = docx.enum.table.WD_CELL_VERTICAL_ALIGNMENT.CENTER

            # Set paragraph style
            p = curr_cell.paragraphs[0]
            p.style = doc.styles["Question"]
            p.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

            # Add text to cell
            add_text_from_str(cell, p, alignment="center")

    # Shade first row if withHeadings is true
    if with_headings:
        for cell in table.rows[0].cells:
            set_table_header_bg_color(cell)


def process_text(text: str) -> str:
    text = text.replace("&nbsp;", " ")
    text = text.replace("\n", " ")  # WARN: Idk if this will break the code but \n usually shouldn't appear
                                    # \n usually appears in the most random places when pasting from Word docx
                                    # EditorJS only represents line breaks as <br> tags
    text = text.replace("<br>", "\n")
    text = text.replace("<b></b>", "")
    text = text.strip("\n")

    return text


def sanitise_mcq_answer(answer: str) -> str:
    """Sanitise the answer for the MCQ question

    Args:
        answer (str): Answer to sanitise

    Returns:
        str: Sanitised answer
    """
    answer = answer.replace("<p>", "")
    answer = answer.replace("</p>", "")
    answer = answer.replace("<br>", "")
    answer = answer.replace("<b>", "")
    answer = answer.replace("</b>", "")
    answer = answer.replace("<i>", "")
    answer = answer.replace("</i>", "")
    answer = answer.replace("<s>", "")
    answer = answer.replace("</s>", "")
    answer = answer.replace("<u>", "")
    answer = answer.replace("</u>", "")
    answer = answer.replace("&nbsp;", "")

    return answer


def insertHR(paragraph: docx.text.paragraph.Paragraph):
    """ Insert a horizontal rule at the end of the paragraph.

    Args:
        paragraph (docx.text.paragraph.Paragraph): Paragraph to insert the horizontal rule into
    """
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
                              'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
                              'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
                              'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
                              'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
                              'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
                              'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
                              'w:pPrChange'
                              )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)


def set_table_header_bg_color(cell):
    tblCell = cell._tc
    tblCellProperties = tblCell.get_or_add_tcPr()
    clShading = OxmlElement('w:shd')
    clShading.set(qn('w:fill'), "D9D9D9")
    tblCellProperties.append(clShading)
    return cell


def delete_paragraph(paragraph):
    """Delete a paragraph from the document.

    Args:
        paragraph (docx.text.paragraph.Paragraph): Paragraph to delete
    """
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None
